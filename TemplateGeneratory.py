from re import S
import tkinter as tk
from tkinter import messagebox, filedialog, ttk
from docx import Document
from docx2pdf import convert
import os
import traceback
import logging
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TemplateGeneratorApp:
    def __init__(self, root):
        # Configure logging
        logging.basicConfig(
            filename='template_generator.log',  # Log file name
            level=logging.DEBUG,                 # Log level (DEBUG to capture all messages)
            format='%(asctime)s - %(levelname)s - %(message)s'  # Log message format
        )
        self.root = root
        self.root.title("Template Generator")
        self.root.geometry("600x600")
        self.template_folder = ""
        self.load_template_folder()
        self.template_mapping = {
            "BuyerWelcome.docx": "Buyer Welcome Letter",
            "CosignmentAgreement.docx": "Consignment Agreement",
            "ConsignmentDisclosure.docx": "Consignment Disclosure",
            "ConsignmentContract.docx": "Consignment Contract",
            "DMVBillofSale.docx": "DMV Bill of Sale",
            "DMVNoticeofSale.docx": "DMV Notice of Sale",
            "DMVStatementofError.docx": "DMV Statement of Error",
            "DMVTitleDelay.docx": "DMV Title Delay",
            "DMVVinCheck.docx": "DMV Vin Check",
            "MunicipalityOwned.docx": "Municipality Owned",
            "PrivacyPolicy.docx": "Privacy Policy",
            "TitleinProcesswithDMV.docx": "Title in Process with DMV",
            "VehicleChecklist.docx": "Vehicle Checklist"
        }
        self.placeholder_mapping = {
    "BuyerWelcome.docx": {
        "{{todays date}}": "The current date.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{year}}": "The year of manufacture.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{date of auction}}": "Date of the auction."
    },
    # Add other templates and their respective placeholders and descriptions...
    # Example:
    "VehicleChecklist.docx": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number."
    },
    "TitleinProcesswithDMV.docx": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{date of auction}}": "Date of the auction.",
        "{{todays date}}": "Todays date."
         },
    "DMVStatementofError.docx": {
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{describe the error or erasure}}": "Please describe what was incorrect on the Title.",
        "{{the entry should read as follows}}": "Please write how the Title should read. If no entry, write 'blank.'",
        "{{seller first and last name}}": "Sellers first and last name.",
        "{{date of auction}}": "Date of the auction."
          },
    "PrivacyPolicy.docx": {
        "{{todays date}}": "Todays date.",
        "{{vin number}}": "The Vehicle Identification Number."
         },
    "MunicipalityOwned.docx": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{date of auction}}": "Date of the auction.",
        "{{todays date}}": "Todays date."
         },
    "DMVVinCheck": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{date of auction}}": "Date of the auction.",
        "{{state name}}": "State that the item/vehicle is titled in.",
        "{{todays date}}": "Todays date."
        },
    "DMVTitleDelay.docx": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{lot number}}": "The assigned Lot Number.",
        "{{todays date}}": "Todays date.",
        "{{buyer first and last name}}": "First and last name of buyer.",
        "{{buyer first name}}": "First name of buyer.",
        "{{buyer address}}": "House number and street name. For example: 123 Main St.",
        "{{buyer city, state, zip}}": "City, State Zip Code. For example: Medford, OR 97504.",
        "{{number of delay days}}": "Number of days the title is expected to be delayed."
        },
    "DMVNoticeofSale.docx": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{contract date}}": "Date of contract.",
        "{{buyer first and last name}}": "First and last name of buyer.",
        "{{buyer address}}": "House number and street name. For example: 123 Main St.",
        "{{buyer city, state, zip}}": "City, State Zip Code. For example: Medford, OR 97504.",
        "{{seller first and last name}}": "First and last name of seller.",
        "{{seller address}}": "House number and stree name. For example: 123 Main St.",
        "{{seller city, state, zip}}": "City, State Zip Code. For example: Medford, OR 97504."
        },
    "DMVBillofSale.docx": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{seller first and last name}}": "First and last name of seller.",
        "{{seller address}}": "House number and stree name. For example: 123 Main St.",
        "{{seller city, state, zip}}": "City, State Zip Code. For example: Medford, OR 97504.",
        "{{release date}}": "Day title will be transferred from seller to buyer."
        },
    "ConsignmentDisclosure.docx": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{todays date}}": "Todays date."
        },
    "ConsignmentContract.docx": {
        "{{seller first and last name}}": "First and last name of seller.",
        "{{seller address}}": "House number and stree name. For example: 123 Main St.",
        "{{seller city, state, zip}}": "City, State Zip Code. For example: Medford, OR 97504.",
        "{{todays date}}": "Todays date.",
        "{{date of auction}}": "Date of the auction.",
        "{{auction name}}": "What is the name of the auction? For example: 2025 Annual Online Auction.",
        "{{auction location}}": "Where is the auction located? For example: The Expo Center.",
        "{{auctioneers name}}": "Name of auctioneer.",
        "{{consignment description}}": "Please describe the item being sold: trailer, car, vehicle, etc.",
        "{{item delivery start date}}": "When do items need to be delivered to the auction site?",
        "{{item delivery end date}}": "When is the last day items need to be delivered to the auction site?",
        "{{commission percentage}}": "What is the commission percentage agreed upon for this sale?",
        "{{trucking dispatch name}}": "What is the name of the trucking dispatch company?",
        "{{vin number}}": "The Vehicle Identification Number."
         },
    "ConsignmentAgreement.docx": {
        "{{year}}": "The year of manufacture.",
        "{{make}}": "The make of the vehicle.",
        "{{model}}": "The model of the vehicle.",
        "{{vin number}}": "The Vehicle Identification Number.",
        "{{todays date}}": "Todays date.",
        "{{seller first and last name}}": "First and last name of seller.",
        "{{seller address}}": "House number and stree name. For example: 123 Main St.",
        "{{seller city, state, zip}}": "City, State Zip Code. For example: Medford, OR 97504."
        },
}  

        self.special_templates = {
        "DMVBillofSale.docx": {"size": "statement", "orientation": "landscape"},
        "DMVStatementofError.docx": {"size": "statement", "orientation": "portrait"},
        "DMVNoticeofSale": {"size": "statement", "orientation": "portrait"}  
            }

        self.button_frame = None

        self.start_screen()
      
    def load_template_folder(self):
        """Load the previously selected template folder from a file."""
        try:
            with open("template_folder.txt", "r") as f:
                self.template_folder = f.read().strip()
                if self.template_folder:
                    messagebox.showinfo("Folder Loaded", f"Template folder loaded: {self.template_folder}")
                    return
        except FileNotFoundError:
            pass  # Handle file not found case

    def start_screen(self):
        self.clear_frame()
        start_frame = tk.Frame(self.root)
        start_frame.pack(pady=20)
        tk.Label(start_frame, text="Welcome to J&C Auction's Template Generator", font=("Arial", 16)).pack(pady=10)
        select_folder_btn = tk.Button(start_frame, text="Select Template Folder", command=self.select_template_folder, bg="#ADD8E6", padx=10, pady=5)
        select_folder_btn.pack(pady=10)

    def select_template_folder(self):
        folder_selected = filedialog.askdirectory()
        if folder_selected:
            self.template_folder = folder_selected
            logging.info(f"Template folder set to: {self.template_folder}")
            messagebox.showinfo("Folder Selected", f"Template folder set to: {self.template_folder}")
            self.save_template_folder()
            self.load_templates()

    def save_template_folder(self):
        # Get the directory where the script is running
        base_dir = os.path.dirname(os.path.abspath(__file__))
        file_path = os.path.join(base_dir, "template_folder.txt")
        with open(file_path, "w") as f:
            f.write(self.template_folder)

    def load_templates(self):
        self.templates = [template for template in self.template_mapping.keys() 
                          if os.path.exists(os.path.join(self.template_folder, template))]
        if not self.templates:
            logging.warning("No templates found in the selected folder.")
            messagebox.showwarning("No Templates Found", "No templates found in the selected folder.")
        else:
            logging.info(f"Loaded templates: {self.templates}")
            self.template_selection_screen()  # Only call this if templates are found

    def template_selection_screen(self):
        self.clear_frame()

        # Main frame
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Configure columns to distribute extra space
        main_frame.columnconfigure(0, weight=1)  # Left margin
        main_frame.columnconfigure(1, weight=0)  # Content column
        main_frame.columnconfigure(2, weight=1)  # Right margin

        # Content frame
        content_frame = tk.Frame(main_frame)
        content_frame.grid(row=0, column=1, sticky="n")

        # Header
        tk.Label(content_frame, text="Select The Templates You'd Like To Generate", 
                 font=("Arial", 20, "bold")).pack(pady=20)

        # Checkbox frame with two columns
        checkbox_frame = tk.Frame(content_frame)
        checkbox_frame.pack(fill=tk.BOTH, expand=True, padx=20)

        # Calculate number of rows needed
        num_templates = len(self.templates)
        num_rows = (num_templates + 1) // 2  # Ceiling division to get rows

        self.check_vars = []
        for i, template in enumerate(self.templates):
            var = tk.BooleanVar()
    
            # Calculate row and column
            row = i // 2
            col = i % 2

            cb = tk.Checkbutton(checkbox_frame, 
                                text=self.template_mapping[template], 
                                variable=var,
                                font=("Arial", 14),  # Larger font
                                indicatoron=1,
                                anchor='w')
    
            # Use grid instead of pack
            cb.grid(row=row, column=col, sticky='w', padx=20, pady=5)
            self.check_vars.append(var)

        # Create button
        create_button = tk.Button(content_frame, 
                                  text="Create Selected Templates", 
                                  command=self.info_input_screen, 
                                  bg="#add8e6", 
                                  font=("Arial", 14),
                                  padx=15, 
                                  pady=10)
        create_button.pack(pady=20)

        # Configure row weights
        main_frame.rowconfigure(0, weight=1)  # Center content vertically


    def show_description(self, placeholder):
        description = ""
        for template in self.placeholder_mapping:
            if placeholder in self.placeholder_mapping[template]:
                description = self.placeholder_mapping[template][placeholder]
                break
        if description:
            messagebox.showinfo("Placeholder Description", description)

    def info_input_screen(self):
        self.clear_frame()
        info_frame = tk.Frame(self.root)
        info_frame.pack(fill=tk.BOTH, expand=True)

        # Configure columns to distribute extra space
        info_frame.grid_columnconfigure(0, weight=1)  # Left margin
        info_frame.grid_columnconfigure(1, weight=0)  # Content column
        info_frame.grid_columnconfigure(2, weight=1)  # Right margin

        # Header
        tk.Label(info_frame, text="Please Complete The Following Information", font=("Arial", 16)).grid(row=0, column=1, pady=10)

        # Content frame (will contain placeholders and buttons)
        content_frame = tk.Frame(info_frame)
        content_frame.grid(row=1, column=1, sticky="n")

        # Placeholder frame
        self.placeholder_frame = ttk.Frame(content_frame)
        self.placeholder_frame.pack(fill=tk.BOTH, expand=True)

        # Button frame
        self.button_frame = tk.Frame(content_frame)
        self.button_frame.pack(pady=10)

        self.back_button = tk.Button(self.button_frame, text="Back", command=self.template_selection_screen, bg="#FF4500", padx=10, pady=5)
        self.back_button.pack(side=tk.LEFT, padx=5)

        self.generate_button = tk.Button(self.button_frame, text="Generate PDFs", command=self.generate_pdf, bg="#add8e6", padx=10, pady=5)
        self.generate_button.pack(side=tk.LEFT, padx=5)

        # Configure row weights
        info_frame.grid_rowconfigure(0, weight=0)  # Header
        info_frame.grid_rowconfigure(1, weight=1)  # Content frame

        self.update_placeholders()

    def update_placeholders(self):
            if not hasattr(self, 'placeholder_frame') or not self.placeholder_frame.winfo_exists():
                self.placeholder_frame = ttk.Frame(self.root)
            self.placeholder_frame.pack(fill=tk.BOTH, expand=True)
            
            if hasattr(self, 'placeholder_frame'):
                for widget in self.placeholder_frame.winfo_children():
                    widget.destroy()
            placeholders = {}
            for i, template in enumerate(self.templates):
                if self.check_vars[i].get():
                    template_placeholders = self.placeholder_mapping.get(template, {})
                    placeholders.update(template_placeholders)
            self.placeholder_entries = {}

            # Create a sub-frame for the grid layout
            grid_frame = ttk.Frame(self.placeholder_frame)
            grid_frame.pack(fill=tk.BOTH, expand=True)

            # Configure columns in the grid frame
            grid_frame.columnconfigure(0, weight=0)  # Label column
            grid_frame.columnconfigure(1, weight=1)  # Entry column
            grid_frame.columnconfigure(2, weight=0)  # Button column

            for row, (placeholder, description) in enumerate(placeholders.items()):
                # Label with larger font
                label = ttk.Label(grid_frame, text=placeholder.upper()[2:-2], font=("Arial", 12))
                label.grid(row=row, column=0, sticky="e", padx=(0, 10), pady=5)

                # Entry with larger font and fixed width
                entry = ttk.Entry(grid_frame, font=("Arial", 12), width=35)
                entry.grid(row=row, column=1, sticky="ew", pady=5)

                # Description button
                description_btn = ttk.Button(grid_frame, text="?", width=3,
                                             command=lambda p=placeholder: self.show_description(p))
                description_btn.grid(row=row, column=2, padx=(10, 0), pady=5)

                self.placeholder_entries[placeholder] = entry

    def process_special_template(self, doc, template_name):
        if template_name in self.special_templates:
            settings = self.special_templates[template_name]
            for section in doc.sections:
                if settings["size"] == "statement":
                    if settings["orientation"] == "landscape":
                        section.page_width = Inches(8.5)
                        section.page_height = Inches(5.5)
                    else:
                        section.page_width = Inches(5.5)
                        section.page_height = Inches(8.5)
            
                # Set narrow margins
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.4)
                section.right_margin = Inches(0.4)

            # Prevent table rows from breaking across pages
            for table in doc.tables:
                table.autofit = False
                table.width = Inches(section.page_width.inches - section.left_margin.inches - section.right_margin.inches)
                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
    def generate_pdf(self):
        try:
            placeholder_values = {}
            vin_number = None

            # Check if placeholder_entries is initialized and not None
            if not hasattr(self, 'placeholder_entries') or self.placeholder_entries is None:
                messagebox.showerror("Error", "Placeholder entries are not set.")
                return

            # Collect placeholder values
            for ph, entry in self.placeholder_entries.items():
                if entry is None:  # Check if entry itself is None
                    messagebox.showerror("Error", f"Entry for {ph} is not set.")
                    return
            
                value = entry.get()  # Get the value from the entry
                placeholder_values[ph] = value
            
                if ph == '{{vin number}}':
                    vin_number = value

            if not vin_number:
                messagebox.showerror("Error", "VIN Number is required.")
                return

            # Prompt for save location each time
            save_folder = filedialog.askdirectory(title="Select Folder to Save PDFs")
            if not save_folder:
                return

            pdf_folder = os.path.join(save_folder, "Generated_PDFs")
            os.makedirs(pdf_folder, exist_ok=True)

            for i, template in enumerate(self.templates):
                if self.check_vars[i].get():
                    try:
                        template_path = os.path.join(self.template_folder, template)
                        if not os.path.exists(template_path):
                            messagebox.showerror("Error", f"Template file not found: {template_path}")
                            continue  # Skip this template and continue with the next one
                    
                        logging.info(f"Processing template: {template}")

                        doc = Document(template_path)  # Attempt to load the document
                        self.process_special_template(doc, template)

                        def replace_text(paragraph):
                            for run in paragraph.runs:
                                for placeholder, value in placeholder_values.items():
                                    if placeholder in run.text:
                                        parts = run.text.split(placeholder)
                                        run.text = parts[0]
                                        for j, part in enumerate(parts[1:]):
                                            if j > 0:
                                                run.text += placeholder
                                            run.text += value + part

                        # Replace text in paragraphs and tables
                        for paragraph in doc.paragraphs:
                            replace_text(paragraph)
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        replace_text(paragraph)

                        temp_doc_path = os.path.join(pdf_folder, f'temp_{template}')
                        doc.save(temp_doc_path)  # Save temporary document
                        logging.info(f"Temporary document saved at: {temp_doc_path}")

                        # Attempt to convert to PDF
                        try:
                            logging.info(f"Attempting to convert {temp_doc_path} to PDF.")
                            convert(temp_doc_path)  # Convert to PDF
                            logging.info(f"Successfully converted document to PDF: {temp_doc_path}")
                        except Exception as convert_error:
                            logging.error(f"Conversion failed for {temp_doc_path}: {str(convert_error)}")
                            logging.error(traceback.format_exc())  # Log full traceback
                            messagebox.showerror("Error", f"Failed to convert document to PDF: {str(convert_error)}")
                            continue  # Skip this template and continue with the next one

                        pdf_file_path = os.path.join(pdf_folder, f"{vin_number}_{template.replace('.docx', '.pdf')}")
                    
                        try:
                            os.rename(temp_doc_path.replace('.docx', '.pdf'), pdf_file_path)
                            os.remove(temp_doc_path)  # Remove temporary document
                        except Exception as rename_error:
                            logging.error(f"Failed to rename or remove temporary document: {str(rename_error)}")
                            messagebox.showerror("Error", f"Failed to rename or remove temporary document: {str(rename_error)}")

                    except Exception as e:
                        logging.error(f"Failed to process template {template}: {str(e)}")
                        messagebox.showerror("Error", f"Failed to process template {template}: {str(e)}")
                        continue  # Continue with the next template

            # Successful PDF Generation
            self.hide_generate_buttons()
            self.show_restart_button()
            messagebox.showinfo("Success", "PDFs generated successfully!")

        except Exception as e:
            error_message = f"An unexpected error occurred: {str(e)}\n\n{traceback.format_exc()}"
            logging.error(error_message)
            messagebox.showerror("Error", error_message)

    def hide_generate_buttons(self):
        if hasattr(self, 'back_button'):
            self.back_button.pack_forget()
        if hasattr(self, 'generate_button'):
            self.generate_button.pack_forget()

    def show_restart_button(self):
        if self.button_frame:
            self.restart_button = tk.Button(self.button_frame, text="Restart", command=self.restart_application, 
                                            bg="#4CAF50", fg="white", padx=10, pady=5)
            self.restart_button.pack(pady=10)

    def restart_application(self):
        if hasattr(self, 'restart_button'):
            self.restart_button.pack_forget()
        self.template_selection_screen() 

    def clear_frame(self):
        for widget in self.root.winfo_children():
            widget.destroy()

if __name__ == "__main__":
    root = tk.Tk()
    app = TemplateGeneratorApp(root)
    root.mainloop()

